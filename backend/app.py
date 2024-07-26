# Updated code considering edge cases
import json
import os
import shutil
from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import pandas as pd
import cv2
from docx import Document
from docx.shared import Inches, Pt
from concurrent.futures import ThreadPoolExecutor
from comtypes.client import CreateObject
import pythoncom
import threading
from datetime import datetime

# Initialize Flask app and CORS
app = Flask(__name__)
CORS(app)

# Configure paths and folders
UPLOAD_FOLDER = 'uploads'
GEN_FOLDER = 'gens'
CHECKPOINT_FILE = 'checkpoint.json'
TEMPLATE_PATH = 'certificate-template.png'

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GEN_FOLDER'] = GEN_FOLDER

# Ensure the upload and generation folders exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
if not os.path.exists(GEN_FOLDER):
    os.makedirs(GEN_FOLDER)

# Status variable and lock for thread safety
status = {"message": ""}
status_lock = threading.Lock()

# Helper function to delete files in a folder
def delete_files_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')

def load_checkpoint():
    if os.path.exists(CHECKPOINT_FILE):
        with open(CHECKPOINT_FILE, 'r') as f:
            return json.load(f)
    else:
        return {"processed_seat_numbers": [], "step": "start"}

def save_checkpoint(seat_number, step):
    checkpoint = load_checkpoint()
    if seat_number != 0:
        checkpoint["processed_seat_numbers"].append(seat_number)
    checkpoint["step"] = step
    
    with open(CHECKPOINT_FILE, 'w') as f:
        json.dump(checkpoint, f)

def remove_checkpoint():
    if os.path.exists(CHECKPOINT_FILE):
        os.remove(CHECKPOINT_FILE)

def load_excel_files(request):
    ms6_file = request.files['ms6File']
    bms_file = request.files['bmsFile']
    
    ms6_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'MS6.xlsx')
    bms_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'BMSOG.xlsx')

    ms6_file.save(ms6_file_path)
    bms_file.save(bms_file_path)

    try:
        df1 = pd.read_excel(ms6_file_path)
        df2 = pd.read_excel(bms_file_path)
    except Exception as e:
        print(f"Error reading Excel files: {e}")
        raise e

    return df1, df2, ms6_file_path, bms_file_path

def process_dataframes(df1, df2):
    df2['FREM'].fillna('null', inplace=True)
    df2['RES'].fillna('null', inplace=True)

    dataT = df2[(df2['RSLT'] == 'P') & (df2['FREM'] == 'null') & (df2['RES'] == 'null')]
    dataT.loc[:, 'Gender'] = dataT['SEX'].apply(lambda x: 'MALE' if x == 1 else 'FEMALE' if x == 2 else 'N/A')
    dataT = pd.merge(dataT, df1[['COLL_NO']], on='COLL_NO', how='left')  # Include CGPA in the merge
    dataT = dataT.sort_values(by='COLL_NO', ascending=True)
    dataT['pno'] = dataT.groupby('COLL_NO').cumcount() + 1
    dataT['COLL_NO'] = dataT['COLL_NO'].apply(lambda x: str(x).zfill(4))
    dataT['pno'] = dataT['pno'].apply(lambda x: str(x).zfill(4))

    return dataT

@app.route('/status', methods=['GET'])
def get_status():
    with status_lock:
        return jsonify(status)

@app.route('/delete-files', methods=['POST'])
def delete_files():
    try:
        delete_files_in_folder(app.config['GEN_FOLDER'])
        delete_files_in_folder(app.config['UPLOAD_FOLDER'])
        remove_checkpoint()  # Remove checkpoint on delete
        return "Files deleted successfully", 200
    except Exception as e:
        print(f"Error deleting files: {e}")
        return "Error deleting files", 500

@app.route('/generate-certificates', methods=['POST'])
def generate_certificates():
    global status

    # Check if checkpoint.json exists and load the checkpoint
    checkpoint = load_checkpoint()
    processed_seat_numbers = checkpoint["processed_seat_numbers"]
    current_step = checkpoint["step"]

    # Retrieve values from checkpoint
    year = request.form['year']
    course_name = request.form['courseName']
    semester_number = request.form['semester']

    if current_step == "certificate_generation":
        # If the process was in the middle of certificate generation
        with status_lock:
            status['message'] = "Resuming certificate generation..."
        
        # Proceed with certificate generation for unprocessed seat numbers
        df1, df2, ms6_file_path, bms_file_path = load_excel_files(request)  # Define this helper function for loading Excel files
        dataT = process_dataframes(df1, df2)  # Define this helper function for processing the dataframes
        output_word_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.docx"))
        output_pdf_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.pdf"))

        try:
            with ThreadPoolExecutor(max_workers=10) as executor:
                futures = [
                    executor.submit(generate_certificate, row.copy(), year, course_name, semester_number)
                    for _, row in dataT.iterrows() if row['SEAT_NO'] not in processed_seat_numbers
                ]
                for future in futures:
                    try:
                        result = future.result()
                        save_checkpoint(result['SEAT_NO'], "certificate_generation")  # Save progress after each successful certificate generation
                    except Exception as e:
                        print(f"Error generating certificate: {e}")
                        with status_lock:
                            status['message'] = f"Error generating certificate: {e}"
                        return "Error generating certificates", 500

            with status_lock:
                status['message'] = "Generating Word document..."
            save_checkpoint(0, "word_document_generation")  # Update checkpoint to indicate Word document generation is starting

            create_word_document(dataT, output_word_path, executor)

            if not os.path.exists(output_word_path):
                raise FileNotFoundError(f"Word document not found at {output_word_path}")

            save_checkpoint(0, "word_document_generated")  # Update checkpoint to indicate Word document generation is complete

            with status_lock:
                status['message'] = "Generating PDF file..."
            save_checkpoint(0, "pdf_generation")  # Update checkpoint to indicate PDF generation is starting

            convert_to_pdf(output_word_path, output_pdf_path)

            response = send_file(output_pdf_path, as_attachment=True, download_name="certificates.pdf")

            with status_lock:
                status['message'] = "Completed"
            remove_checkpoint()  # Remove checkpoint on success

        except Exception as e:
            print(f"Error generating certificates: {e}")
            with status_lock:
                status['message'] = "Error generating certificates"
            return "Error generating certificates", 500

        finally:
            # Clean up temporary files
            os.remove(ms6_file_path)
            os.remove(bms_file_path)

        return response

    elif current_step == "word_document_generation":
        # If the process was in the middle of Word document generation
        with status_lock:
            status['message'] = "Resuming Word document generation..."

        output_word_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.docx"))
        output_pdf_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.pdf"))

        try:
            df1, df2, ms6_file_path, bms_file_path = load_excel_files(request)  # Define this helper function for loading Excel files
            dataT = process_dataframes(df1, df2)  # Define this helper function for processing the dataframes
            create_word_document(dataT, output_word_path, ThreadPoolExecutor(max_workers=10))

            if not os.path.exists(output_word_path):
                raise FileNotFoundError(f"Word document not found at {output_word_path}")

            save_checkpoint(0, "word_document_generated")  # Update checkpoint to indicate Word document generation is complete

            with status_lock:
                status['message'] = "Generating PDF file..."
            save_checkpoint(0, "pdf_generation")  # Update checkpoint to indicate PDF generation is starting

            convert_to_pdf(output_word_path, output_pdf_path)

            response = send_file(output_pdf_path, as_attachment=True, download_name="certificates.pdf")

            with status_lock:
                status['message'] = "Completed"
            remove_checkpoint()  # Remove checkpoint on success

        except Exception as e:
            print(f"Error generating certificates: {e}")
            with status_lock:
                status['message'] = "Error generating certificates"
            return "Error generating certificates", 500

        finally:
            # Clean up temporary files
            os.remove(ms6_file_path)
            os.remove(bms_file_path)

        return response

    # If no checkpoint or completed checkpoint, start from the beginning
    if 'ms6File' not in request.files or 'bmsFile' not in request.files or 'year' not in request.form or 'courseName' not in request.form or 'semester' not in request.form:
        return "No file part or year, course name, or semester", 400

    ms6_file = request.files['ms6File']
    bms_file = request.files['bmsFile']
    year = request.form['year']
    course_name = request.form['courseName']
    semester_number = request.form['semester']

    ms6_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'MS6.xlsx')
    bms_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'BMS.xlsx')

    ms6_file.save(ms6_file_path)
    bms_file.save(bms_file_path)

    try:
        df1 = pd.read_excel(ms6_file_path)
        df2 = pd.read_excel(bms_file_path)
    except Exception as e:
        print(f"Error reading Excel files: {e}")
        return "Error reading Excel files", 500

    df2['FREM'].fillna('null', inplace=True)
    df2['RES'].fillna('null', inplace=True)

    dataT = df2[(df2['RSLT'] == 'P') & (df2['FREM'] == 'null') & (df2['RES'] == 'null')]
    dataT.loc[:, 'Gender'] = dataT['SEX'].apply(lambda x: 'MALE' if x == 1 else 'FEMALE' if x == 2 else 'N/A')
    dataT = pd.merge(dataT, df1[['COLL_NO']], on='COLL_NO', how='left')  # Include CGPA in the merge
    dataT = dataT.sort_values(by='COLL_NO', ascending=True)
    dataT['pno'] = dataT.groupby('COLL_NO').cumcount() + 1
    dataT['COLL_NO'] = dataT['COLL_NO'].apply(lambda x: str(x).zfill(4))
    dataT['pno'] = dataT['pno'].apply(lambda x: str(x).zfill(4))

    output_word_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.docx"))
    output_pdf_path = os.path.abspath(os.path.join(app.config['GEN_FOLDER'], "certificates.pdf"))

    try:
        with status_lock:
            status['message'] = "Generating certificates..."

        with ThreadPoolExecutor(max_workers=10) as executor:
            futures = [
                executor.submit(generate_certificate, row.copy(), year, course_name, semester_number)
                for _, row in dataT.iterrows() if row['SEAT_NO'] not in processed_seat_numbers
            ]
            for future in futures:
                try:
                    result = future.result()
                    save_checkpoint(result['SEAT_NO'], "certificate_generation")  # Save progress after each successful certificate generation
                except Exception as e:
                    print(f"Error generating certificate: {e}")
                    with status_lock:
                        status['message'] = f"Error generating certificate: {e}"
                    return "Error generating certificates", 500

            with status_lock:
                status['message'] = "Generating Word document..."
            save_checkpoint(0, "word_document_generation")  # Update checkpoint to indicate Word document generation is starting

            create_word_document(dataT, output_word_path, executor)

            if not os.path.exists(output_word_path):
                raise FileNotFoundError(f"Word document not found at {output_word_path}")

            save_checkpoint(0, "word_document_generated")  # Update checkpoint to indicate Word document generation is complete

            with status_lock:
                status['message'] = "Generating PDF file..."
            save_checkpoint(0, "pdf_generation")  # Update checkpoint to indicate PDF generation is starting

            convert_to_pdf(output_word_path, output_pdf_path)

            response = send_file(output_pdf_path, as_attachment=True, download_name="certificates.pdf")

            with status_lock:
                status['message'] = "Completed"
            remove_checkpoint()  # Remove checkpoint on success

    except Exception as e:
        print(f"Error generating certificates: {e}")
        with status_lock:
            status['message'] = "Error generating certificates"
        return "Error generating certificates", 500

    finally:
        # Clean up temporary files
        os.remove(ms6_file_path)
        os.remove(bms_file_path)

    return response

def generate_certificate(row, year, course_name, semester_number):
    certificate_template_image = cv2.imread(TEMPLATE_PATH)
    
    name = str(row['NAME']).strip() if pd.notnull(row['NAME']) else 'N/A'
    coll_no = str(row['COLL_NO'])
    pno = str(row['pno'])
    ccf = f"CCF : {coll_no} : {pno}"
    seat_no = str(row['SEAT_NO']).strip() if pd.notnull(row['SEAT_NO']) else 'N/A'
    seat_no1 = "NO : " + seat_no
    gender = "/ - FEMALE" if pd.notnull(row['Gender']) and row['Gender'] == 'FEMALE' else ''
    name_with_gender = f"/ {name}" if gender else name
    semester_roman = convert_to_roman(int(semester_number))
    thirdline = "held by the University of Mumbai in the month of"
    director = "DIRECTOR"
    board = "BOARD OF EXAMINATIONS & EVALUATION"

    current_date = datetime.now().strftime("%B %d, %Y")  # Get the current date in the format "MONTH DATE, YEAR"

    if 'CGPA' in row:
        cgpa = str(row['CGPA']) if pd.notnull(row['CGPA']) else 'N/A'
        course_semester_text = f"PASSED THE {course_name} (SEM {semester_roman}) (CBCGS) EXAMINATION"
        date_text = f"{year} WITH {cgpa} CGPI"
    elif 'GRADE' in row:
        grade = str(row['GRADE']) if pd.notnull(row['GRADE']) else 'N/A'
        course_semester_text = f"PASSED THE {course_name} (SEM {semester_roman}) (CBSGS) EXAMINATION"
        date_text = f"{year} AND WAS PLACED IN THE {grade} GRADE"
    # else:
    #     date_text = year

    filename = f"{seat_no}.png"
    output_path = os.path.join(app.config['GEN_FOLDER'], filename)

    try:
        cv2.putText(certificate_template_image, ccf, (300, 590), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, seat_no1, (300, 680), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, name_with_gender, (400, 1300), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, course_semester_text, (385, 1500), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, thirdline, (385, 1700), cv2.FONT_HERSHEY_COMPLEX, 2, (0, 0, 0), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, date_text, (390, 1850), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, gender, (350, 2300), cv2.FONT_HERSHEY_COMPLEX, 2, (157, 157, 157), 4, cv2.LINE_AA)  
        cv2.putText(certificate_template_image, current_date, (330, 2400), cv2.FONT_HERSHEY_COMPLEX, 2,(157, 157, 157), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, director, (1700, 2300), cv2.FONT_HERSHEY_COMPLEX, 2,(0, 0, 0), 4, cv2.LINE_AA)
        cv2.putText(certificate_template_image, board, (1200, 2400), cv2.FONT_HERSHEY_COMPLEX, 2,(0, 0, 0), 4, cv2.LINE_AA)

        cv2.imwrite(output_path, certificate_template_image)
        print(f"Generated certificate: {output_path}")

        return {"SEAT_NO": seat_no}

    except Exception as e:
        print(f"Error generating certificate for SEAT_NO {seat_no}: {e}")
        raise

def convert_to_roman(number):
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syb = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_numeral = ''
    for i in range(len(val)):
        count = int(number / val[i])
        roman_numeral += syb[i] * count
        number -= val[i] * count
    return roman_numeral

def create_word_document(data, output_word_path, executor):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0)

    count = 0  # Counter to keep track of certificates added

    for i, row in data.iterrows():
        seat_no = str(row['SEAT_NO']).strip() if pd.notnull(row['SEAT_NO']) else 'N/A'
        name = str(row['NAME']).strip() if pd.notnull(row['NAME']) else 'N/A'
        img_path = os.path.join(app.config['GEN_FOLDER'], f"{seat_no}.png")
        if os.path.exists(img_path):
            # Add the certificate image
            # doc.add_picture(img_path, width=Inches(7), height=Inches(4.6))
            doc.add_picture(img_path, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # Center alignment for the image
            
            # Add a 0.3-inch space after the image
            space_paragraph = doc.add_paragraph()
            space_run = space_paragraph.add_run()
            space_run.add_break()
            space_paragraph_format = space_paragraph.paragraph_format
            space_paragraph_format.space_before = Pt(0)
            space_paragraph_format.space_after = Pt(0)

            count += 1  # Increment the counter

            # Add a page break after every 2 certificates
            if count % 2 != 0 and i < len(data) - 1:
                space_paragraph_format.line_spacing = Pt(3.6 * 5)  # 0.3 inch space

    doc.save(output_word_path)

def convert_to_pdf(input_word_path, output_pdf_path):
    pythoncom.CoInitialize()
    word = CreateObject("Word.Application")
    doc = word.Documents.Open(input_word_path)
    doc.SaveAs(output_pdf_path, FileFormat=17)
    doc.Close()
    word.Quit()
    pythoncom.CoUninitialize()

if __name__ == '__main__':
    app.run(debug=True)
